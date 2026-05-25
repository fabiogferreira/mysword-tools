"""
Cria um documento Word de exemplo para testar a conversão para Journal MySword.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def create_sample_docx():
    """Cria um documento Word de exemplo com estudos bíblicos."""
    
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ================================================
    # TÍTULO PRINCIPAL
    # ================================================
    title = doc.add_heading('Estudo: O Sermão do Monte', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    intro = doc.add_paragraph()
    intro.add_run('Um estudo sobre os ensinamentos de Jesus registrados em Mateus 5-7').italic = True
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Espaço
    
    # ================================================
    # SEÇÃO 1: INTRODUÇÃO
    # ================================================
    doc.add_heading('Introdução ao Sermão do Monte', level=1)
    
    p1 = doc.add_paragraph()
    p1.add_run('O Sermão do Monte é considerado o mais importante discurso de Jesus Cristo. ')
    p1.add_run('Encontrado em Mateus 5-7').bold = True
    p1.add_run(', este sermão contém os princípios fundamentais do Reino de Deus.')
    
    doc.add_paragraph(
        'Jesus subiu ao monte e, vendo as multidões, assentou-se. '
        'Seus discípulos aproximaram-se dele, e ele começou a ensiná-los.'
    )
    
    # Lista de tópicos
    doc.add_paragraph('O sermão aborda:', style='List Bullet')
    doc.add_paragraph('As Bem-aventuranças (Mateus 5:3-12)', style='List Bullet')
    doc.add_paragraph('Sal da terra e luz do mundo (Mateus 5:13-16)', style='List Bullet')
    doc.add_paragraph('A Lei e os Profetas (Mateus 5:17-48)', style='List Bullet')
    doc.add_paragraph('Oração e jejum (Mateus 6:1-18)', style='List Bullet')
    doc.add_paragraph('Tesouros no céu (Mateus 6:19-34)', style='List Bullet')
    doc.add_paragraph('Não julgueis (Mateus 7:1-6)', style='List Bullet')
    doc.add_paragraph('Pedi e será dado (Mateus 7:7-12)', style='List Bullet')
    
    # ================================================
    # SEÇÃO 2: AS BEM-AVENTURANÇAS
    # ================================================
    doc.add_heading('As Bem-aventuranças', level=1)
    
    doc.add_paragraph(
        'As Bem-aventuranças são nove declarações de bênçãos pronunciadas por Jesus '
        'no início do Sermão do Monte. Cada uma começa com "Bem-aventurados..." '
        '(em grego: μακάριοι - makarioi).'
    )
    
    # Subtítulo
    doc.add_heading('Texto Bíblico (Mateus 5:3-12)', level=2)
    
    # Citação bíblica com formatação
    verses = [
        (3, 'Bem-aventurados os pobres de espírito, porque deles é o Reino dos céus.'),
        (4, 'Bem-aventurados os que choram, porque serão consolados.'),
        (5, 'Bem-aventurados os mansos, porque herdarão a terra.'),
        (6, 'Bem-aventurados os que têm fome e sede de justiça, porque serão fartos.'),
        (7, 'Bem-aventurados os misericordiosos, porque alcançarão misericórdia.'),
        (8, 'Bem-aventurados os limpos de coração, porque verão a Deus.'),
        (9, 'Bem-aventurados os pacificadores, porque serão chamados filhos de Deus.'),
        (10, 'Bem-aventurados os perseguidos por causa da justiça, porque deles é o Reino dos céus.'),
    ]
    
    for verse_num, verse_text in verses:
        p = doc.add_paragraph()
        run_num = p.add_run(f'{verse_num} ')
        run_num.bold = True
        run_num.font.color.rgb = RGBColor(0, 100, 180)
        p.add_run(verse_text)
    
    # ================================================
    # SEÇÃO 3: SAL E LUZ
    # ================================================
    doc.add_heading('Sal da Terra e Luz do Mundo', level=1)
    
    doc.add_paragraph(
        'Após as Bem-aventuranças, Jesus usa duas metáforas poderosas para descrever '
        'a missão dos seus discípulos no mundo.'
    )
    
    # Subtítulo
    doc.add_heading('Sal da Terra (Mateus 5:13)', level=2)
    
    quote1 = doc.add_paragraph()
    quote1.add_run(
        '"Vós sois o sal da terra; e, se o sal for insípido, com que se há de salgar? '
        'Para nada mais presta, senão para se lançar fora e ser pisado pelos homens."'
    ).italic = True
    
    doc.add_paragraph(
        'O sal era usado para preservar alimentos e dar sabor. '
        'Da mesma forma, os cristãos devem preservar os valores do Reino '
        'e trazer "sabor" à sociedade através de seu testemunho.'
    )
    
    # Subtítulo
    doc.add_heading('Luz do Mundo (Mateus 5:14-16)', level=2)
    
    quote2 = doc.add_paragraph()
    quote2.add_run(
        '"Vós sois a luz do mundo; não se pode esconder uma cidade edificada sobre um monte."'
    ).italic = True
    
    doc.add_paragraph(
        'A luz representa a verdade e a presença de Deus. '
        'Os discípulos são chamados a iluminar um mundo em trevas, '
        'não escondendo sua fé, mas deixando-a brilhar para que outros vejam '
        'e glorifiquem a Deus.'
    )
    
    # ================================================
    # SEÇÃO 4: A ORAÇÃO DO PAI NOSSO
    # ================================================
    doc.add_heading('A Oração do Pai Nosso', level=1)
    
    doc.add_paragraph(
        'Em Mateus 6:9-13, Jesus ensina seus discípulos como orar, '
        'dando-lhes um modelo de oração conhecido como "Pai Nosso" ou "Oração do Senhor".'
    )
    
    # Tabela com a estrutura da oração
    doc.add_heading('Estrutura da Oração', level=2)
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    # Cabeçalho
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Parte da Oração'
    header_cells[1].text = 'Significado'
    
    # Dados
    data = [
        ('Pai nosso que estás nos céus', 'Relacionamento íntimo com Deus'),
        ('Santificado seja o teu nome', 'Adoração e reverência'),
        ('Venha o teu Reino', 'Submissão à vontade de Deus'),
        ('O pão nosso de cada dia', 'Dependência diária de Deus'),
        ('Perdoa as nossas dívidas', 'Confissão e perdão'),
        ('Livra-nos do mal', 'Proteção espiritual'),
    ]
    
    for i, (parte, significado) in enumerate(data, 1):
        row_cells = table.rows[i].cells
        row_cells[0].text = parte
        row_cells[1].text = significado
    
    # ================================================
    # SEÇÃO 5: APLICAÇÃO
    # ================================================
    doc.add_heading('Aplicação Prática', level=1)
    
    doc.add_paragraph(
        'O Sermão do Monte não é apenas um conjunto de ideais elevados, '
        'mas um chamado à transformação prática de vida. Considere estas aplicações:'
    )
    
    # Lista numerada
    applications = [
        'Examine seu coração à luz das Bem-aventuranças',
        'Seja sal e luz em seu ambiente de trabalho e família',
        'Pratique a oração do Pai Nosso diariamente',
        'Busque primeiro o Reino de Deus (Mateus 6:33)',
        'Construa sua vida sobre a rocha - as palavras de Jesus (Mateus 7:24-27)',
    ]
    
    for i, app in enumerate(applications, 1):
        doc.add_paragraph(f'{i}. {app}')
    
    # Versículo final em destaque
    doc.add_paragraph()
    
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.add_run(
        '"Buscai primeiro o Reino de Deus e a sua justiça, '
        'e todas estas coisas vos serão acrescentadas."'
    ).bold = True
    
    ref = doc.add_paragraph('— Mateus 6:33')
    ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salvar documento
    output_dir = Path(__file__).parent.parent / "examples"
    output_dir.mkdir(exist_ok=True)
    
    output_file = output_dir / "sermao_do_monte.docx"
    doc.save(str(output_file))
    
    print(f"✅ Documento Word criado: {output_file}")
    return output_file


if __name__ == "__main__":
    create_sample_docx()
