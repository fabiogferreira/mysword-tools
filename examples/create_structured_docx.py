"""
Cria um documento Word de exemplo com metadados estruturados para MySword Tools.
Demonstra como estruturar um documento para extração automática de metadados.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_structured_docx():
    """
    Cria um documento Word com metadados estruturados.
    
    Estrutura demonstrada:
    1. Propriedades do documento (File > Properties)
    2. Metadados inline no início do documento
    3. Títulos Heading 1 para separar entradas
    """
    
    doc = Document()
    
    # === PROPRIEDADES DO DOCUMENTO (serão lidas automaticamente) ===
    doc.core_properties.title = "Estudos sobre Oração"
    doc.core_properties.author = "Pastor João Silva"
    doc.core_properties.subject = "Um guia prático sobre como desenvolver uma vida de oração eficaz"
    doc.core_properties.keywords = "oração, vida cristã, espiritualidade"
    
    # === TÍTULO PRINCIPAL ===
    title = doc.add_heading('Estudos sobre Oração', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    subtitle = doc.add_paragraph()
    subtitle.add_run('Um guia prático para desenvolver uma vida de oração').italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # === METADADOS INLINE (alternativa às propriedades do documento) ===
    # Estes são lidos se as propriedades do documento estiverem vazias
    # Comentado aqui porque já definimos nas propriedades
    # doc.add_paragraph('Autor: Pastor João Silva')
    # doc.add_paragraph('Descrição: Um guia prático sobre oração')
    # doc.add_paragraph('Tags: oração, vida cristã, espiritualidade')
    # doc.add_paragraph('Abreviação: ORACAO')
    
    doc.add_paragraph()
    
    # ================================================
    # ENTRADA 1: Introdução
    # ================================================
    doc.add_heading('Introdução à Oração', level=1)
    
    doc.add_paragraph(
        'A oração é o meio pelo qual nos comunicamos com Deus. '
        'Através dela, expressamos gratidão, fazemos pedidos, confessamos '
        'nossos pecados e buscamos direção para nossas vidas.'
    )
    
    doc.add_paragraph(
        'Jesus nos ensinou a importância da oração em Mateus 6:6: '
        '"Mas tu, quando orares, entra no teu quarto e, fechando a porta, '
        'ora a teu Pai que está em secreto..."'
    )
    
    # Lista de benefícios
    doc.add_paragraph('Benefícios da Oração:', style='List Bullet')
    doc.add_paragraph('Desenvolve intimidade com Deus', style='List Bullet')
    doc.add_paragraph('Traz paz e direção', style='List Bullet')
    doc.add_paragraph('Fortalece a fé', style='List Bullet')
    doc.add_paragraph('Transforma o caráter', style='List Bullet')
    
    # ================================================
    # ENTRADA 2: Tipos de Oração
    # ================================================
    doc.add_heading('Tipos de Oração', level=1)
    
    doc.add_paragraph(
        'A Bíblia nos mostra diferentes tipos de oração, cada uma com '
        'seu propósito específico na vida do crente.'
    )
    
    # Subtítulo
    doc.add_heading('1. Oração de Adoração', level=2)
    doc.add_paragraph(
        'É quando exaltamos a Deus por quem Ele é, não pelo que Ele faz. '
        'Exemplo: Salmo 103:1 - "Bendize, ó minha alma, ao Senhor, '
        'e tudo o que há em mim bendiga o seu santo nome."'
    )
    
    doc.add_heading('2. Oração de Confissão', level=2)
    doc.add_paragraph(
        'Reconhecemos nossos pecados diante de Deus e pedimos perdão. '
        '1 João 1:9 nos garante: "Se confessarmos os nossos pecados, '
        'ele é fiel e justo para nos perdoar..."'
    )
    
    doc.add_heading('3. Oração de Súplica', level=2)
    doc.add_paragraph(
        'Apresentamos nossas necessidades a Deus. Filipenses 4:6 '
        'nos encoraja: "Não andeis ansiosos por coisa alguma; '
        'em tudo, porém, sejam conhecidas as vossas petições..."'
    )
    
    doc.add_heading('4. Oração Intercessória', level=2)
    doc.add_paragraph(
        'Oramos pelos outros. Paulo frequentemente intercedia pelas igrejas, '
        'como vemos em Efésios 1:16-18.'
    )
    
    # ================================================
    # ENTRADA 3: Modelo de Oração de Jesus
    # ================================================
    doc.add_heading('O Pai Nosso - Modelo de Oração', level=1)
    
    doc.add_paragraph(
        'Em Mateus 6:9-13, Jesus nos ensinou o modelo perfeito de oração, '
        'conhecido como "Pai Nosso".'
    )
    
    # Tabela com estrutura do Pai Nosso
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    header = table.rows[0].cells
    header[0].text = 'Parte da Oração'
    header[1].text = 'Significado'
    
    data = [
        ('Pai nosso que estás nos céus', 'Reconhecimento de Deus como Pai'),
        ('Santificado seja o teu nome', 'Adoração e reverência'),
        ('Venha o teu Reino', 'Submissão à vontade de Deus'),
        ('O pão nosso de cada dia', 'Dependência diária'),
        ('Perdoa as nossas dívidas', 'Confissão e perdão'),
        ('Livra-nos do mal', 'Proteção espiritual'),
    ]
    
    for i, (parte, significado) in enumerate(data, 1):
        row = table.rows[i].cells
        row[0].text = parte
        row[1].text = significado
    
    doc.add_paragraph()
    
    # ================================================
    # ENTRADA 4: Prática Diária
    # ================================================
    doc.add_heading('Desenvolvendo uma Vida de Oração', level=1)
    
    doc.add_paragraph(
        'Para desenvolver uma vida de oração consistente, considere estas práticas:'
    )
    
    practices = [
        'Estabeleça um horário fixo para oração',
        'Escolha um lugar tranquilo sem distrações',
        'Use um diário de oração para registrar pedidos e respostas',
        'Comece com gratidão antes de fazer pedidos',
        'Inclua leitura bíblica em seu tempo de oração',
        'Ore com outros crentes regularmente',
    ]
    
    for i, practice in enumerate(practices, 1):
        doc.add_paragraph(f'{i}. {practice}')
    
    # Versículo final
    doc.add_paragraph()
    final = doc.add_paragraph()
    final.add_run(
        '"Orai sem cessar." - 1 Tessalonicenses 5:17'
    ).bold = True
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Salvar
    output_dir = Path(__file__).parent
    output_file = output_dir / "estudo_oracao.docx"
    doc.save(str(output_file))
    
    print(f"✅ Documento criado: {output_file}")
    print()
    print("Metadados do documento:")
    print(f"  Título: {doc.core_properties.title}")
    print(f"  Autor: {doc.core_properties.author}")
    print(f"  Assunto: {doc.core_properties.subject}")
    print(f"  Palavras-chave: {doc.core_properties.keywords}")
    
    return output_file


if __name__ == "__main__":
    create_structured_docx()
