"""
Cria um documento Word modelo completo demonstrando todas as opções
de metadados e estrutura para o MySword Tools.

Este documento serve como template/exemplo para criação de seus próprios estudos.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def create_template_docx():
    """
    Cria um documento Word template completo com:
    1. Propriedades do documento configuradas
    2. Metadados inline demonstrados
    3. Estrutura completa de entradas
    4. Exemplos de formatação suportada
    """
    
    doc = Document()
    
    # =====================================================
    # PROPRIEDADES DO DOCUMENTO (METADADOS AUTOMÁTICOS)
    # =====================================================
    # Estes campos são lidos automaticamente pelo conversor:
    doc.core_properties.title = "Template de Estudo Bíblico"
    doc.core_properties.author = "Seu Nome Aqui"
    doc.core_properties.subject = "Um modelo completo para criar estudos bíblicos compatíveis com MySword"
    doc.core_properties.keywords = "template, modelo, estudo bíblico, mysword"
    doc.core_properties.comments = "Criado com MySword Tools"
    
    # =====================================================
    # CABEÇALHO DO DOCUMENTO
    # =====================================================
    
    # Título principal (estilo "Title")
    title = doc.add_heading('Template de Estudo Bíblico', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo
    subtitle = doc.add_paragraph()
    run = subtitle.add_run('Um modelo completo para criar estudos compatíveis com MySword')
    run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # =====================================================
    # METADADOS INLINE (ALTERNATIVA ÀS PROPRIEDADES)
    # =====================================================
    # Se as propriedades do documento estiverem vazias,
    # estes padrões serão usados:
    
    meta_section = doc.add_paragraph()
    meta_section.add_run('─' * 50)
    
    # Exemplo de metadados inline (comentados para não duplicar)
    meta_info = doc.add_paragraph()
    meta_info.add_run('METADADOS DO DOCUMENTO\n').bold = True
    meta_info.add_run('(Estes campos são lidos automaticamente pelo conversor)\n\n')
    meta_info.add_run('• ').bold = True
    meta_info.add_run('Autor: ').bold = True
    meta_info.add_run('Definido em Arquivo > Propriedades > Autor\n')
    meta_info.add_run('• ').bold = True
    meta_info.add_run('Descrição: ').bold = True
    meta_info.add_run('Definido em Arquivo > Propriedades > Assunto\n')
    meta_info.add_run('• ').bold = True
    meta_info.add_run('Tags: ').bold = True
    meta_info.add_run('Definido em Arquivo > Propriedades > Palavras-chave\n')
    meta_info.add_run('\n')
    meta_info.add_run('Ou adicione no início do documento:\n').italic = True
    meta_info.add_run('  Autor: Seu Nome\n')
    meta_info.add_run('  Descrição: Descrição do estudo\n')
    meta_info.add_run('  Tags: tag1, tag2, tag3\n')
    meta_info.add_run('  Abreviação: ABREV\n')
    
    meta_end = doc.add_paragraph()
    meta_end.add_run('─' * 50)
    
    doc.add_paragraph()
    
    # =====================================================
    # ENTRADA 1: COMO USAR ESTE TEMPLATE
    # =====================================================
    doc.add_heading('Como Usar Este Template', level=1)
    
    doc.add_paragraph(
        'Este documento demonstra a estrutura recomendada para criar estudos '
        'bíblicos que serão convertidos para o formato MySword Journal.'
    )
    
    doc.add_heading('Estrutura Básica', level=2)
    
    doc.add_paragraph(
        'O conversor divide o documento em entradas baseado nos títulos. '
        'Cada "Título 1" (Heading 1) cria uma nova entrada no Journal.'
    )
    
    # Lista de regras
    rules = [
        'Use Título 1 (Heading 1) para separar entradas do Journal',
        'Use Título 2, 3, etc para subdivisões dentro de cada entrada',
        'Parágrafos normais formam o conteúdo de cada entrada',
        'Tabelas são convertidas para HTML automaticamente',
        'Formatação (negrito, itálico, cores) é preservada',
        'Referências bíblicas são convertidas em links clicáveis',
    ]
    
    for rule in rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    # =====================================================
    # ENTRADA 2: FORMATAÇÃO SUPORTADA
    # =====================================================
    doc.add_heading('Formatação Suportada', level=1)
    
    doc.add_paragraph('O conversor preserva diversos tipos de formatação:')
    
    doc.add_heading('Estilos de Texto', level=2)
    
    p = doc.add_paragraph()
    p.add_run('Texto em ')
    run_bold = p.add_run('negrito')
    run_bold.bold = True
    p.add_run(', ')
    run_italic = p.add_run('itálico')
    run_italic.italic = True
    p.add_run(', ')
    run_under = p.add_run('sublinhado')
    run_under.underline = True
    p.add_run(' e ')
    run_color = p.add_run('colorido')
    run_color.font.color.rgb = RGBColor(0, 102, 204)
    p.add_run('.')
    
    doc.add_heading('Listas', level=2)
    
    doc.add_paragraph('Lista com marcadores:')
    doc.add_paragraph('Primeiro item', style='List Bullet')
    doc.add_paragraph('Segundo item', style='List Bullet')
    doc.add_paragraph('Terceiro item', style='List Bullet')
    
    doc.add_paragraph('Lista numerada:')
    doc.add_paragraph('Passo um', style='List Number')
    doc.add_paragraph('Passo dois', style='List Number')
    doc.add_paragraph('Passo três', style='List Number')
    
    doc.add_heading('Citações', level=2)
    
    quote = doc.add_paragraph()
    quote.add_run(
        '"Lâmpada para os meus pés é a tua palavra, e luz para o meu caminho."'
    ).italic = True
    quote.add_run('\n— Salmo 119:105')
    
    # =====================================================
    # ENTRADA 3: REFERÊNCIAS BÍBLICAS
    # =====================================================
    doc.add_heading('Referências Bíblicas', level=1)
    
    doc.add_paragraph(
        'Referências bíblicas no texto são automaticamente convertidas em links '
        'clicáveis no MySword. Por exemplo:'
    )
    
    refs_items = [
        'João 3:16 - Versículo único',
        'Mateus 5:3-12 - Intervalo de versículos',
        'Gênesis 1:1 - Livro por extenso',
        'Gn 1:1 - Livro abreviado',
        'Sl 23:1 - Salmos',
        'Rm 8:28 - Romanos',
        '1 Co 13:4-8 - Livros com número',
    ]
    
    for item in refs_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('Formatos de Referência Aceitos', level=2)
    
    # Tabela de formatos
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    
    headers = ['Formato', 'Exemplo', 'Resultado no MySword']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    data = [
        ('Livro Cap:Vers', 'João 3:16', 'Link para João 3:16'),
        ('Livro Cap:Vers-Vers', 'Mt 5:1-12', 'Link para intervalo'),
        ('Abreviação padrão', 'Gn, Ex, Lv...', 'Reconhece abreviações'),
        ('Com espaço', '1 Co 13:4', 'Livros numerados'),
        ('Português', 'Gênesis 1:1', 'Nomes em português'),
    ]
    
    for i, (fmt, ex, result) in enumerate(data, 1):
        table.rows[i].cells[0].text = fmt
        table.rows[i].cells[1].text = ex
        table.rows[i].cells[2].text = result
    
    doc.add_paragraph()
    
    # =====================================================
    # ENTRADA 4: TABELAS
    # =====================================================
    doc.add_heading('Usando Tabelas', level=1)
    
    doc.add_paragraph(
        'Tabelas do Word são convertidas para HTML e exibidas corretamente '
        'no MySword. Use tabelas para organizar informações.'
    )
    
    doc.add_heading('Exemplo: Frutos do Espírito', level=2)
    
    table2 = doc.add_table(rows=10, cols=3)
    table2.style = 'Table Grid'
    
    headers2 = ['Fruto', 'Grego', 'Referência']
    for i, h in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    
    fruits = [
        ('Amor', 'ἀγάπη (agapē)', 'Gálatas 5:22'),
        ('Alegria', 'χαρά (chara)', 'Gálatas 5:22'),
        ('Paz', 'εἰρήνη (eirēnē)', 'Gálatas 5:22'),
        ('Paciência', 'μακροθυμία (makrothymia)', 'Gálatas 5:22'),
        ('Benignidade', 'χρηστότης (chrēstotēs)', 'Gálatas 5:22'),
        ('Bondade', 'ἀγαθωσύνη (agathōsynē)', 'Gálatas 5:22'),
        ('Fidelidade', 'πίστις (pistis)', 'Gálatas 5:22'),
        ('Mansidão', 'πραΰτης (prautēs)', 'Gálatas 5:23'),
        ('Domínio próprio', 'ἐγκράτεια (enkrateia)', 'Gálatas 5:23'),
    ]
    
    for i, (fruto, grego, ref) in enumerate(fruits, 1):
        table2.rows[i].cells[0].text = fruto
        table2.rows[i].cells[1].text = grego
        table2.rows[i].cells[2].text = ref
    
    doc.add_paragraph()
    
    # =====================================================
    # ENTRADA 5: DICAS FINAIS
    # =====================================================
    doc.add_heading('Dicas para Criação de Estudos', level=1)
    
    doc.add_heading('Boas Práticas', level=2)
    
    tips = [
        'Preencha as propriedades do documento para metadados automáticos',
        'Use Título 1 para cada tópico principal do estudo',
        'Mantenha parágrafos concisos para melhor leitura em dispositivos móveis',
        'Inclua referências bíblicas no formato padrão para links automáticos',
        'Teste o resultado no MySword após a conversão',
    ]
    
    for i, tip in enumerate(tips, 1):
        doc.add_paragraph(f'{i}. {tip}')
    
    doc.add_heading('Conversão via Linha de Comando', level=2)
    
    doc.add_paragraph('Comando básico (usa metadados do documento):')
    cmd1 = doc.add_paragraph()
    cmd1.add_run('python -m src.word_to_journal documento.docx saida.jor.mybible').italic = True
    
    doc.add_paragraph('Com parâmetros personalizados:')
    cmd2 = doc.add_paragraph()
    cmd2.add_run(
        'python -m src.word_to_journal documento.docx saida.jor.mybible '
        '--title "Meu Título" --author "Meu Nome"'
    ).italic = True
    
    doc.add_paragraph()
    
    # Versículo final
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    final.add_run(
        '"Toda Escritura é inspirada por Deus e útil para o ensino, '
        'para a repreensão, para a correção, para a educação na justiça."'
    ).bold = True
    
    ref_final = doc.add_paragraph('— 2 Timóteo 3:16')
    ref_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # =====================================================
    # SALVAR DOCUMENTO
    # =====================================================
    output_dir = Path(__file__).parent
    output_file = output_dir / "template_estudo_biblico.docx"
    doc.save(str(output_file))
    
    print(f"✅ Template criado: {output_file}")
    print()
    print("📋 Propriedades do documento:")
    print(f"   Título: {doc.core_properties.title}")
    print(f"   Autor: {doc.core_properties.author}")
    print(f"   Assunto: {doc.core_properties.subject}")
    print(f"   Palavras-chave: {doc.core_properties.keywords}")
    print()
    print("📝 Estrutura do documento:")
    print("   - Como Usar Este Template")
    print("   - Formatação Suportada")
    print("   - Referências Bíblicas")
    print("   - Usando Tabelas")
    print("   - Dicas para Criação de Estudos")
    
    return output_file


if __name__ == "__main__":
    create_template_docx()
