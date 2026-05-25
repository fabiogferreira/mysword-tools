"""
Script de teste para validar a extração de imagens inline em Base64
"""
import os
import sys
import sqlite3
import base64
from io import BytesIO
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Inches

# Adiciona o diretório raiz ao path do Python
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src.word_to_journal import WordToJournalConverter

def create_image_test_docx(filepath):
    doc = Document()
    
    doc.add_paragraph("Autor: Teste de Imagem")
    doc.add_paragraph("Abreviação: IMGTST")
    doc.add_paragraph("")
    
    # Adicionar seção de Introdução
    doc.add_paragraph("@@---@@")
    doc.add_paragraph("Lição com Imagem")
    doc.add_paragraph("Abaixo está uma imagem dinâmica azul criada para validar o suporte a imagens no backend:")
    
    # Criar uma imagem simples azul 100x100 em memória
    img = Image.new('RGB', (100, 100), color='blue')
    img_io = BytesIO()
    img.save(img_io, format='PNG')
    img_io.seek(0)
    
    # Inserir imagem no documento Word
    doc.add_picture(img_io, width=Inches(1.5))
    
    doc.add_paragraph("Fim do documento com imagem.")
    doc.save(filepath)
    print(f"Documento de teste com imagem criado em: {filepath}")

def test_image_conversion():
    input_docx = "examples/test_image_doc.docx"
    output_db = "output/test_image_doc.jor.mybible"
    
    os.makedirs("examples", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    
    create_image_test_docx(input_docx)
    
    if os.path.exists(output_db):
        os.remove(output_db)
        
    print(f"\nConvertendo o documento: {input_docx}")
    converter = WordToJournalConverter(input_docx)
    journal = converter.convert(split_by_heading=True)
    journal.save(output_db)
    
    # Validar se o Base64 da imagem foi parar na tabela 'journal' do SQLite
    print("\nValidando banco SQLite gerado...")
    conn = sqlite3.connect(output_db)
    cursor = conn.cursor()
    
    cursor.execute("SELECT title, content FROM journal WHERE title LIKE '%Imagem%'")
    row = cursor.fetchone()
    
    if row:
        title, content = row
        print(f"  Encontrada seção: '{title}'")
        
        # Procura a substring de base64 no HTML gerado
        if "data:image/png;base64," in content:
            print("  SUCCESS: Tag <img src=\"data:image/png;base64,...\" /> encontrada no conteudo!")
            # Imprime uma amostra do conteúdo contendo a tag
            amostra_start = content.find("data:image/png;base64,")
            print(f"  Amostra da tag: {content[amostra_start-30:amostra_start+50]}...")
        else:
            print("  FAIL: A string Base64 da imagem nao foi encontrada no conteudo.")
    else:
        print("  FAIL: Nao foi encontrada a secao com a imagem no banco.")
        
    conn.close()

if __name__ == "__main__":
    test_image_conversion()
