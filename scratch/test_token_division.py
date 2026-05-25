"""
Script de teste para validar a divisão de seções baseada no token @@---@@
"""
import os
import sys
import sqlite3
from pathlib import Path
from docx import Document

# Adiciona o diretório raiz ao path do Python
sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from src.word_to_journal import WordToJournalConverter

def create_test_docx(filepath):
    doc = Document()
    
    # Adicionar metadados no início
    doc.add_paragraph("Autor: Autor de Teste")
    doc.add_paragraph("Descrição: Estudo testando divisão por token")
    doc.add_paragraph("Abreviação: TESTTOK")
    doc.add_paragraph("")
    
    # Seção Inicial / Introdução
    doc.add_paragraph("Conteúdo introdutório do documento que não pertence a nenhuma lição específica.")
    doc.add_paragraph("")
    
    # Primeira Lição
    doc.add_paragraph("@@---@@")
    doc.add_paragraph("Lição 1: O Começo de Tudo")
    doc.add_paragraph("Este é o conteúdo da lição 1. Referência: Gênesis 1:1.")
    doc.add_paragraph("")
    
    # Segunda Lição
    doc.add_paragraph("@@---@@")
    doc.add_paragraph("Lição 2: O Caminho no Deserto")
    doc.add_paragraph("Este é o conteúdo da lição 2. Referência: Êxodo 14:14.")
    doc.add_paragraph("")
    
    # Terceira Lição (Vazia de conteúdo, mas com título)
    doc.add_paragraph("@@---@@")
    doc.add_paragraph("Lição 3: Conclusão")
    doc.add_paragraph("")
    
    doc.save(filepath)
    print(f"Documento de teste criado: {filepath}")

def test_token_division():
    input_docx = "examples/test_token.docx"
    output_db = "output/test_token.jor.mybible"
    
    os.makedirs("examples", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    
    create_test_docx(input_docx)
    
    if os.path.exists(output_db):
        os.remove(output_db)
        
    print(f"\nAnalisando arquivo: {input_docx}")
    converter = WordToJournalConverter(input_docx)
    
    # Executa a crítica (análise de estrutura)
    suggestions = converter.critique(split_by_heading=True, heading_level=1)
    print("\nCrítica / Sugestões:")
    for sug in suggestions:
        print(f"  [{sug.level}] {sug.message}")
        
    # Executa a conversão
    print(f"\nConvertendo para: {output_db}")
    journal = converter.convert(
        split_by_heading=True,
        heading_level=1
    )
    journal.save(output_db)
    
    # Valida as tabelas geradas no SQLite
    print("\nValidando banco de dados gerado...")
    conn = sqlite3.connect(output_db)
    cursor = conn.cursor()
    
    cursor.execute("SELECT COUNT(*) FROM journal")
    count = cursor.fetchone()[0]
    print(f"  Número total de entradas na tabela 'journal': {count}")
    
    cursor.execute("SELECT title, content FROM journal")
    rows = cursor.fetchall()
    print("  Entradas geradas:")
    for idx, (title, content) in enumerate(rows, 1):
        print(f"    {idx}. {title}")
        print(f"       Conteúdo: {content.strip()}")
        
    conn.close()

if __name__ == "__main__":
    test_token_division()
