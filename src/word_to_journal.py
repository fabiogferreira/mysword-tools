"""
Word to Journal Converter - Converte documentos Word (.docx) para Journal MySword (.jor.mybible)

Este módulo processa documentos Word e os converte para o formato Journal do MySword,
preservando formatação, estilos e estrutura do documento.
"""

import re
import html
from pathlib import Path
from typing import List, Tuple, Optional
from dataclasses import dataclass, field

@dataclass
class CritiqueSuggestion:
    """Uma sugestão de melhoria ou aviso de formatação"""
    level: str  # 'INFO', 'WARNING', 'ERROR'
    message: str
    suggestion: str = ""

class HTMLListManager:
    """Gerencia a abertura e fechamento de tags <ul> e <ol>."""
    def __init__(self, get_list_tag_func):
        self.stack = []  # Lista de (num_id, ilvl, tag)
        self.get_list_tag = get_list_tag_func

    def handle_paragraph(self, info: dict) -> str:
        html_out = []
        is_list = info['is_list']
        new_ilvl = info['list_lvl']
        new_num_id = info['list_id']

        if not is_list:
            # Fecha todas as listas abertas
            while self.stack:
                _, _, tag = self.stack.pop()
                html_out.append(f"</{tag}>")
            return "".join(html_out)

        # Determina a tag se for uma nova lista
        new_tag = self.get_list_tag(new_num_id, new_ilvl)

        # Caso 1: Nível aumentou ou primeira lista
        if not self.stack or new_ilvl > self.stack[-1][1]:
            while not self.stack or new_ilvl > self.stack[-1][1]:
                lvl_to_add = len(self.stack)
                self.stack.append((new_num_id, lvl_to_add, new_tag))
                html_out.append(f"<{new_tag}>")
        
        # Caso 2: Nível diminuiu
        elif new_ilvl < self.stack[-1][1]:
            while self.stack and new_ilvl < self.stack[-1][1]:
                _, _, tag = self.stack.pop()
                html_out.append(f"</{tag}>")
            
            # Se mudou o numId ou o tipo de lista no mesmo nível
            if self.stack and (self.stack[-1][0] != new_num_id or self.stack[-1][2] != new_tag):
                _, _, tag = self.stack.pop()
                html_out.append(f"</{tag}>")
                self.stack.append((new_num_id, new_ilvl, new_tag))
                html_out.append(f"<{new_tag}>")

        # Caso 3: Mesmo nível mas mudou a lista (numId)
        elif self.stack[-1][0] != new_num_id:
             _, _, tag = self.stack.pop()
             html_out.append(f"</{tag}>")
             self.stack.append((new_num_id, new_ilvl, new_tag))
             html_out.append(f"<{new_tag}>")

        return "".join(html_out)

    def close_all(self) -> str:
        html_out = []
        while self.stack:
            _, _, tag = self.stack.pop()
            html_out.append(f"</{tag}>")
        return "".join(html_out)

from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .mysword_journal import MySwordJournal, JournalEntry


@dataclass
class ConversionOptions:
    """Opções de conversão Word → Journal"""
    # Metadados do Journal
    abbreviation: str
    title: str = ""  # Título exibido no MySword
    description: str = ""
    author: str = ""  # Autor do documento
    readonly: bool = False
    
    # Opções de conversão
    split_by_heading: bool = True  # Dividir documento por títulos
    heading_level: int = 1  # Nível do heading para dividir (1=Título 1, 2=Título 2)
    preserve_formatting: bool = True  # Preservar negrito, itálico, etc.
    include_images: bool = False  # Incluir imagens (como base64) - futuro
    convert_tables: bool = True  # Converter tabelas para HTML
    
    # CSS personalizado
    customcss: str = ""


class WordToJournalConverter:
    """
    Converte documentos Word (.docx) para Journal do MySword (.jor.mybible).
    
    O conversor pode:
    - Dividir o documento em múltiplas entradas baseado em títulos
    - Preservar formatação (negrito, itálico, sublinhado, cores)
    - Converter tabelas para HTML
    - Detectar referências bíblicas e converter para links
    
    Uso:
        converter = WordToJournalConverter("documento.docx")
        journal = converter.convert(
            abbreviation="MEU_DOC",
            title="Meu Documento",
            description="Documento convertido"
        )
        journal.save("documento.jor.mybible")
    """
    
    # Padrão para detectar referências bíblicas
    BIBLE_REF_PATTERN = re.compile(
        r'\b((?:'
        r'Gênesis|Genesis|Gên|Gen|Gn|'
        r'Êxodo|Exodo|Êx|Ex|'
        r'Levítico|Levitico|Lev|Lv|'
        r'Números|Numeros|Núm|Num|Nm|'
        r'Deuteronômio|Deuteronomio|Deut|Dt|'
        r'Josué|Josue|Js|'
        r'Juízes|Juizes|Juí|Jui|Jz|'
        r'Rute|Rut|Rt|'
        r'1\s*Samuel|1\s*Sam|1\s*Sm|'
        r'2\s*Samuel|2\s*Sam|2\s*Sm|'
        r'1\s*Reis|1\s*Rs|1\s*Re|'
        r'2\s*Reis|2\s*Rs|2\s*Re|'
        r'1\s*Crônicas|1\s*Cronicas|1\s*Cr|'
        r'2\s*Crônicas|2\s*Cronicas|2\s*Cr|'
        r'Esdras|Esd|Ed|'
        r'Neemias|Nee|Ne|'
        r'Ester|Et|'
        r'Provérbios|Proverbios|Prov|Pv|'
        r'Eclesiastes|Ec|Ecl|'
        r'Cânticos|Canticos|Cantares|Cant|Ct|'
        r'Isaías|Isaias|Isa|Is|'
        r'Lamentações|Lamentacoes|Lam|Lm|'
        r'Ezequiel|Ez|'
        r'Daniel|Dn|Dan|'
        r'Oséias|Oseias|Osé|Ose|Os|'
        r'Joel|Jl|Joe|'
        r'Amós|Amos|Amó|Amo|Am|'
        r'Obadias|Ob|Oba|'
        r'Jonas|Jn|Jon|'
        r'Miquéias|Miqueias|Miq|Mq|'
        r'Naum|Na|'
        r'Habacuque|Hab|Hc|'
        r'Sofonias|Sof|Sf|'
        r'Ageu|Ag|Hag|'
        r'Zacarias|Zac|Zc|'
        r'Malaquias|Mal|Ml|'
        r'Mateus|Mat|Mt|'
        r'Marcos|Mar|Mc|'
        r'Lucas|Luc|Lc|'
        r'João|Joao|John|Joh|Jo|'  # Colocado Jo aqui e removido do Job
        r'Atos|Ato|At|'
        r'Romanos|Rom|Rm|'
        r'1\s*Coríntios|1\s*Corintios|1\s*Cor|1\s*Co|'
        r'2\s*Coríntios|2\s*Corintios|2\s*Cor|2\s*Co|'
        r'Gálatas|Galatas|Gál|Gal|Gl|'
        r'Efésios|Efesios|Efé|Efe|Ef|'
        r'Filipenses|Fil|Fp|'
        r'Colossenses|Col|Cl|'
        r'1\s*Tessalonicenses|1\s*Tes|1\s*Ts|'
        r'2\s*Tessalonicenses|2\s*Tes|2\s*Ts|'
        r'1\s*Timóteo|1\s*Timoteo|1\s*Tim|1\s*Tm|'
        r'2\s*Timóteo|2\s*Timoteo|2\s*Tim|2\s*Tm|'
        r'Tito|Tit|Tt|'
        r'Filemom|Flm|Fm|'
        r'Hebreus|Heb|Hb|'
        r'Tiago|Tia|Tg|'
        r'1\s*Pedro|1\s*Ped|1\s*Pe|'
        r'2\s*Pedro|2\s*Ped|2\s*Pe|'
        r'1\s*João|1\s*Joao|1\s*Joh|1\s*Jo|'
        r'2\s*João|2\s*Joao|2\s*Joh|2\s*Jo|'
        r'3\s*João|3\s*Joao|3\s*Joh|3\s*Jo|'
        r'Judas|Jud|Jd|'
        r'Apocalipse|Apo|Ap|'
        r'Salmos|Salmo|Sal|Sl|'
        r'Jeremias|Jr|'
        r'Jó|Job'  # Jó/Job sem Jo para evitar conflito com João
        r'))\s*'
        r'(\d{1,3})'  # Capítulo
        r'[:\.]'  # Separador
        r'(\d{1,3})'  # Versículo inicial
        r'(?:[-–](\d{1,3}))?'  # Versículo final (opcional)
        r'\b',
        re.IGNORECASE
    )
    
    def __init__(self, filepath: str):
        """
        Inicializa o conversor com um arquivo Word.
        
        Args:
            filepath: Caminho para o arquivo .docx
        """
        self.filepath = Path(filepath)
        if not self.filepath.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {filepath}")
        if self.filepath.suffix.lower() != '.docx':
            raise ValueError("O arquivo deve ter extensão .docx")
        
        self.document: DocumentType = Document(str(self.filepath))
        
        # Extrai metadados do documento
        self._extracted_metadata = self._extract_metadata()
    
    def _extract_metadata(self) -> dict:
        """
        Extrai metadados do documento Word de múltiplas fontes:
        1. Propriedades do documento (File > Properties)
        2. Padrões no texto do documento (Autor:, Descrição:, Tags:)
        3. Primeiro título como nome do Journal
        
        Returns:
            Dicionário com metadados extraídos
        """
        metadata = {
            'title': '',
            'author': '',
            'description': '',
            'tags': '',
            'abbreviation': ''
        }
        
        # 1. Extrai das propriedades do documento Word
        core_props = self.document.core_properties
        if core_props.title:
            metadata['title'] = core_props.title
        if core_props.author:
            metadata['author'] = core_props.author
        if core_props.subject:
            metadata['description'] = core_props.subject
        if core_props.keywords:
            metadata['tags'] = core_props.keywords
        
        # 2. Procura padrões no início do documento
        # Padrões suportados:
        #   Autor: Nome do Autor
        #   Descrição: Descrição do documento
        #   Tags: tag1, tag2, tag3
        #   Abreviação: ABREV
        metadata_patterns = {
            'author': re.compile(r'^(?:Autor|Author)\s*:\s*(.+)$', re.IGNORECASE),
            'description': re.compile(r'^(?:Descrição|Description|Desc)\s*:\s*(.+)$', re.IGNORECASE),
            'tags': re.compile(r'^(?:Tags?|Palavras[- ]?chave|Keywords?)\s*:\s*(.+)$', re.IGNORECASE),
            'abbreviation': re.compile(r'^(?:Abreviação|Abbreviation|Abrev)\s*:\s*(.+)$', re.IGNORECASE),
        }
        
        # Processa os primeiros parágrafos (antes do segundo heading)
        heading_count = 0
        first_heading_text = ""
        
        for para in self.document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Verifica estilo inline (para evitar dependência de método ainda não definido)
            style_name = para.style.name if para.style else ""
            is_heading = style_name.startswith('Heading') or style_name.startswith('Título') or style_name == 'Title'
            
            # Captura o primeiro heading como título potencial
            if is_heading:
                heading_count += 1
                if heading_count == 1:
                    first_heading_text = text
                if heading_count > 1:
                    # Para de procurar após o segundo heading
                    break
            
            # Procura padrões de metadados em parágrafos normais
            if not is_heading and heading_count <= 1:
                for key, pattern in metadata_patterns.items():
                    match = pattern.match(text)
                    if match:
                        # Só sobrescreve se ainda não tiver valor
                        if not metadata[key]:
                            metadata[key] = match.group(1).strip()
        
        # Se não temos título, usa o primeiro heading
        if not metadata['title'] and first_heading_text:
            metadata['title'] = first_heading_text
        
        # Se não temos título, usa o nome do arquivo
        if not metadata['title']:
            metadata['title'] = self.filepath.stem
        
        # Gera abreviação se não tiver
        if not metadata['abbreviation']:
            # Cria abreviação a partir do título
            abbr = metadata['title'].upper().replace(' ', '_')[:20]
            abbr = re.sub(r'[^A-Z0-9_]', '', abbr)
            metadata['abbreviation'] = abbr or self.filepath.stem.upper()[:20]
        
        return metadata

    def _has_division_tokens(self) -> bool:
        """
        Verifica se o documento contém o token de divisão explícita @@---@@.
        """
        for p in self.document.paragraphs:
            if p.text.strip() == "@@---@@":
                return True
        return False

    def critique(self, split_by_heading: bool = True, heading_level: int = 1) -> List[CritiqueSuggestion]:
        """
        Analisa o documento e gera uma lista de críticas e sugestões.
        
        Args:
            split_by_heading: Se o documento será dividido
            heading_level: Nível de heading esperado para divisão
            
        Returns:
            Lista de CritiqueSuggestion
        """
        suggestions = []
        metadata = self._extracted_metadata
        
        # 1. Validação de Metadados
        if not metadata['author']:
            suggestions.append(CritiqueSuggestion(
                'INFO', 
                "Autor não encontrado nas propriedades ou no texto.",
                "Adicione 'Autor: Seu Nome' no início do documento."
            ))
        
        if len(metadata['abbreviation']) > 20:
            suggestions.append(CritiqueSuggestion(
                'WARNING',
                f"Abreviação '{metadata['abbreviation']}' é muito longa.",
                "Use uma abreviação de no máximo 20 caracteres."
            ))

        # 2. Análise de Estrutura
        has_token = self._has_division_tokens()
        
        if has_token:
            token_count = sum(1 for p in self.document.paragraphs if p.text.strip() == "@@---@@")
            suggestions.append(CritiqueSuggestion(
                'INFO',
                f"O documento contém {token_count} token(s) de divisão explícita '@@---@@'.",
                "A divisão de seções será feita utilizando esses tokens."
            ))
        elif split_by_heading:
            headings = []
            total_p = 0
            long_sections = 0
            current_section_length = 0
            
            for para in self.document.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                    
                total_p += 1
                style_type, level = self._get_paragraph_style_type(para)
                
                if style_type == 'heading':
                    headings.append((text, level))
                    if current_section_length > 50: # Arbitrário: mais de 50 parágrafos por seção
                        long_sections += 1
                    current_section_length = 0
                else:
                    current_section_length += 1

            # Tenta detecção se heading_level for padrão ou 0
            actual_level = heading_level if heading_level > 0 else self._detect_best_heading_level()
            
            headings = []
            for para in self.document.paragraphs:
                style_type, level = self._get_paragraph_style_type(para)
                if style_type == 'heading':
                    headings.append((para.text.strip(), level))
            
            if not headings:
                suggestions.append(CritiqueSuggestion(
                    'WARNING',
                    "Nenhum título (Heading) encontrado para dividir o documento.",
                    "Aplique estilos de 'Título' (Heading 1, 2, etc) nas seções ou use o token @@---@@."
                ))
            else:
                target_headings = [h for h in headings if h[1] == actual_level]
                if not target_headings:
                    # Tenta nível superior
                    target_headings = [h for h in headings if h[1] <= actual_level]
                
                num_entries = len(target_headings) if target_headings else 1
                suggestions.append(CritiqueSuggestion(
                    'INFO',
                    f"O documento será dividido em {num_entries} tópico(s) (blocos internos).",
                    f"Nível de divisão detectado: Título {actual_level}."
                ))
                
                if long_sections > 0:
                    suggestions.append(CritiqueSuggestion(
                        'INFO',
                        f"Encontradas {long_sections} seções muito longas.",
                        "Considere adicionar mais títulos para facilitar a navegação no MySword."
                    ))

        # 3. Referências Bíblicas
        bible_refs = 0
        for para in self.document.paragraphs:
            if self.BIBLE_REF_PATTERN.search(para.text):
                bible_refs += len(self.BIBLE_REF_PATTERN.findall(para.text))
        
        if bible_refs == 0:
            suggestions.append(CritiqueSuggestion(
                'INFO',
                "Nenhuma referência bíblica detectada.",
                "Verifique se o formato das referências está correto (ex: João 3:16)."
            ))
        elif bible_refs > 0:
            suggestions.append(CritiqueSuggestion(
                'INFO',
                f"Detectadas {bible_refs} referências bíblicas que serão transformadas em links.",
                ""
            ))

        return suggestions
    
    def get_extracted_metadata(self) -> dict:
        """
        Retorna os metadados extraídos do documento.
        
        Returns:
            Dicionário com: title, author, description, tags, abbreviation
        """
        return self._extracted_metadata.copy()

    def _detect_best_heading_level(self) -> int:
        """
        Analisa o documento para encontrar o melhor nível de heading para dividir.
        Busca o nível que tem mais ocorrências e/ou contém palavras-chave como 'Aula'.
        
        Returns:
            Nível do heading (1-6)
        """
        level_counts = {}
        level_keywords = {} # Níveis que contêm palavras como 'Aula', 'Lição'
        
        for para in self.document.paragraphs:
            style_type, level = self._get_paragraph_style_type(para)
            if style_type == 'heading':
                level_counts[level] = level_counts.get(level, 0) + 1
                text = para.text.lower()
                if 'aula' in text or 'lição' in text or 'licao' in text:
                    level_keywords[level] = level_keywords.get(level, 0) + 1
        
        if not level_counts:
            return 1
            
        # Prioridade 1: Nível com palavras-chave
        if level_keywords:
            return max(level_keywords, key=level_keywords.get)
            
        # Prioridade 2: Nível mais frequente entre 1 e 3
        filtered_counts = {l: c for l, c in level_counts.items() if l <= 3}
        if filtered_counts:
            return max(filtered_counts, key=filtered_counts.get)
            
        return max(level_counts, key=level_counts.get)

    def _generate_smart_id(self, title: str, current_month: Optional[int] = None) -> str:
        """
        Gera um ID inteligente baseado no título (ex: 'Aula 1' -> 'aula1').
        Tenta incluir o mês se disponível.
        """
        text = title.lower()
        # Tenta encontrar número da aula/lição
        lesson_match = re.search(r'(?:aula|lição|licao)\s*(\d+)', text)
        lesson_num = lesson_match.group(1) if lesson_match else ""
        
        if lesson_num:
            prefix = f"mes{current_month}" if current_month else ""
            return f"{prefix}aula{lesson_num}"
            
        # Fallback para slug normal
        clean = re.sub(r'[^a-z0-9]', '', text)
        return clean[:30]
    
    def _get_paragraph_info(self, paragraph: Paragraph) -> dict:
        """
        Extrai informações detalhadas sobre o estilo e propriedades do parágrafo.
        
        Returns:
            Dicionário com: type, level, is_list, list_id, list_lvl
        """
        style_name = paragraph.style.name if paragraph.style else ""
        info = {
            'type': 'normal',
            'level': 0,
            'is_list': False,
            'list_id': None,
            'list_lvl': 0
        }
        
        # Detecta headings
        if style_name.startswith('Heading') or style_name.startswith('Título'):
            match = re.search(r'(\d+)', style_name)
            info['type'] = 'heading'
            info['level'] = int(match.group(1)) if match else 1
            return info
            
        # Detecta propriedades de lista no XML (mais robusto que apenas nome do estilo)
        p_pr = paragraph._p.pPr
        if p_pr is not None and p_pr.numPr is not None:
            info['is_list'] = True
            info['type'] = 'list'
            if p_pr.numPr.numId is not None:
                info['list_id'] = p_pr.numPr.numId.val
            if p_pr.numPr.ilvl is not None:
                info['list_lvl'] = p_pr.numPr.ilvl.val
            return info
            
        # Fallback para nomes de estilo
        if 'List' in style_name or 'Lista' in style_name:
            info['is_list'] = True
            info['type'] = 'list'
            return info
            
        if 'Quote' in style_name or 'Citação' in style_name:
            info['type'] = 'quote'
            return info
            
        return info

    def _get_paragraph_style_type(self, paragraph: Paragraph) -> Tuple[str, int]:
        """
        Identifica o tipo de estilo do parágrafo (Mantido para retrocompatibilidade).
        """
        info = self._get_paragraph_info(paragraph)
        return (info['type'], info['level'])
    
    def _run_to_html(self, run: Run, preserve_formatting: bool = True) -> str:
        """
        Converte um Run do Word para HTML, suportando imagens inline convertidas para base64.
        
        Args:
            run: Objeto Run do python-docx
            preserve_formatting: Se True, preserva formatação
            
        Returns:
            String HTML
        """
        text = html.escape(run.text)
        
        # Verifica se o run contém desenhos de imagem
        img_htmls = []
        try:
            drawings = run.element.xpath('.//w:drawing')
            for drawing in drawings:
                blips = drawing.xpath('.//*[local-name()="blip"]')
                if blips:
                    rId = blips[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if rId:
                        # Busca a imagem correspondente
                        image_part = run.part.related_parts[rId]
                        image_bytes = image_part.image.blob
                        content_type = image_part.image.content_type
                        
                        import base64
                        encoded = base64.b64encode(image_bytes).decode('utf-8')
                        
                        # Gera tag img responsiva
                        img_htmls.append(f'<div style="text-align:center;margin:12px 0;"><img src="data:{content_type};base64,{encoded}" style="max-width:100%;height:auto;border-radius:6px;box-shadow:0 2px 8px rgba(0,0,0,0.15);" /></div>')
        except Exception:
            # Silencia erros se falhar no XML de imagens
            pass
            
        if not text and not img_htmls:
            return ""
            
        if text and preserve_formatting:
            # Aplica formatação
            if run.bold:
                text = f"<b>{text}</b>"
            if run.italic:
                text = f"<i>{text}</i>"
            if run.underline:
                text = f"<u>{text}</u>"
            
            # Cor do texto
            if run.font.color and run.font.color.rgb:
                color = str(run.font.color.rgb)
                text = f"<span style='color:#{color}'>{text}</span>"
            
            # Texto sobrescrito ou subscrito
            if run.font.superscript:
                text = f"<sup>{text}</sup>"
            if run.font.subscript:
                text = f"<sub>{text}</sub>"
        
        # Retorna o texto formatado seguido de qualquer imagem contida no run
        return text + "".join(img_htmls)
    
    def _paragraph_to_html(self, paragraph: Paragraph, 
                           preserve_formatting: bool = True,
                           convert_bible_refs: bool = True) -> str:
        """
        Converte um parágrafo do Word para HTML (apenas conteúdo interno para listas).
        """
        # Coleta o conteúdo HTML dos runs
        html_parts = []
        for run in paragraph.runs:
            html_parts.append(self._run_to_html(run, preserve_formatting))
        
        content = ''.join(html_parts)
        
        if not content.strip() and not paragraph.text.strip():
            return ""
            
        # Converte referências bíblicas para links
        if convert_bible_refs:
            content = self._convert_bible_references(content)
            
        return content

    def _get_list_tag(self, num_id: int, ilvl: int = 0) -> str:
        """
        Tenta descobrir se a lista é ordenada (ol) ou não (ul) para um nível específico.
        """
        if num_id is None:
            return 'ul'
            
        try:
            import lxml.etree as ET
            # Acessa o XML bruto do numbering.xml
            numbering_part = self.document.part.numbering_part
            xml_content = numbering_part.blob
            root = ET.fromstring(xml_content)
            
            # Namespaces necessários para o XPath
            nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # Busca o numId correspondente
            nums = root.xpath(f'//w:num[@w:numId="{num_id}"]', namespaces=nsmap)
            if not nums:
                return 'ul'
                
            num = nums[0]
            abstract_num_ids = num.xpath('.//w:abstractNumId/@w:val', namespaces=nsmap)
            if not abstract_num_ids:
                return 'ul'
                
            abstract_num_id = abstract_num_ids[0]
            abstract_nums = root.xpath(f'//w:abstractNum[@w:abstractNumId="{abstract_num_id}"]', namespaces=nsmap)
            if not abstract_nums:
                return 'ul'
                
            abstract_num = abstract_nums[0]
            
            # Pega o formato do nível solicitado (ilvl)
            fmts = abstract_num.xpath(f'.//w:lvl[@w:ilvl="{ilvl}"]/w:numFmt/@w:val', namespaces=nsmap)
            if not fmts:
                # Fallback para ilvl=0 se não achar o específico
                fmts = abstract_num.xpath('.//w:lvl[@w:ilvl="0"]/w:numFmt/@w:val', namespaces=nsmap)
                
            if fmts:
                fmt = fmts[0]
                if fmt in ['decimal', 'lowerLetter', 'upperLetter', 'lowerRoman', 'upperRoman']:
                    return 'ol'
        except Exception:
            # Fallback seguro diante de qualquer erro no parsing
            pass
            
        return 'ul'
    
    def _table_to_html(self, table: Table) -> str:
        """
        Converte uma tabela do Word para HTML.
        
        Args:
            table: Objeto Table do python-docx
            
        Returns:
            String HTML da tabela
        """
        html_parts = ["<table border='1' cellpadding='5' cellspacing='0'>"]
        
        for row_idx, row in enumerate(table.rows):
            html_parts.append("<tr>")
            for cell in row.cells:
                # Usa <th> para primeira linha (cabeçalho)
                tag = "th" if row_idx == 0 else "td"
                cell_text = html.escape(cell.text)
                html_parts.append(f"<{tag}>{cell_text}</{tag}>")
            html_parts.append("</tr>")
        
        html_parts.append("</table>")
        return '\n'.join(html_parts)
    
    def _convert_bible_references(self, text: str) -> str:
        """
        Converte referências bíblicas em texto para links MySword.
        
        Args:
            text: Texto com possíveis referências bíblicas
            
        Returns:
            Texto com referências convertidas para links
        """
        def replace_ref(match):
            book = match.group(1)
            chapter = match.group(2)
            verse_start = match.group(3)
            verse_end = match.group(4) if match.group(4) else None
            
            full_match = match.group(0)
            
            if verse_end:
                ref = f"{book} {chapter}:{verse_start}-{verse_end}"
            else:
                ref = f"{book} {chapter}:{verse_start}"
            
            return f"<a href='b{ref}'>{full_match}</a>"
        
        return self.BIBLE_REF_PATTERN.sub(replace_ref, text)
    
    def _split_by_headings(self, heading_level: int = 1) -> List[Tuple[str, str, str]]:
        """
        Divide o documento em seções baseado nos headings ou no token @@---@@.
        """
        sections = []
        current_title = self.filepath.stem
        current_content = []
        last_detected_month = None
        current_section_month = None
        
        list_manager = HTMLListManager(self._get_list_tag)
        
        # Otimização O(N): Mapeia element para parágrafo/tabela
        p_map = {p._element: p for p in self.document.paragraphs}
        tbl_map = {t._element: t for t in self.document.tables}
        
        has_token = self._has_division_tokens()
        expecting_title = False
        
        # Detecta nível se necessário
        if not has_token and heading_level <= 0:
            heading_level = self._detect_best_heading_level()
        
        for element in self.document.element.body:
            if element.tag.endswith('p'):
                para = p_map.get(element)
                
                if para:
                    info = self._get_paragraph_info(para)
                    style_type = info['type']
                    level = info['level']
                    text = para.text.strip()
                    
                    if has_token:
                        if text == "@@---@@":
                            # Fecha a seção anterior
                            current_content.append(list_manager.close_all())
                            if current_content:
                                sections.append((
                                    current_title,
                                    '\n'.join(current_content),
                                    self._generate_smart_id(current_title, current_section_month)
                                ))
                            current_content = []
                            expecting_title = True
                            continue
                            
                        if expecting_title:
                            if text:  # Primeiro parágrafo não vazio vira o título da seção
                                current_title = text
                                expecting_title = False
                                current_section_month = last_detected_month
                                continue
                            else:
                                continue
                    
                    if not has_token:
                        # Detecta mês para ID inteligente (Heading 2 no exemplo)
                        if style_type == 'heading' and level < heading_level:
                            month_match = re.search(r'(?:mês|mes)\s*(\d+)', text.lower())
                            if month_match:
                                last_detected_month = int(month_match.group(1))
                        
                        # Se encontrou heading do nível especificado, inicia nova seção
                        if style_type == 'heading' and level == heading_level:
                            # Fecha listas abertas antes de fechar a seção
                            current_content.append(list_manager.close_all())
                            
                            if current_content:
                                sections.append((
                                    current_title,
                                    '\n'.join(current_content),
                                    self._generate_smart_id(current_title, current_section_month)
                                ))
                            
                            current_title = text
                            current_content = []
                            current_section_month = last_detected_month
                            continue
                    
                    # Gerencia tags de lista
                    list_tags_html = list_manager.handle_paragraph(info)
                    if list_tags_html:
                        current_content.append(list_tags_html)
                        
                    # Converte conteúdo do parágrafo
                    inner_html = self._paragraph_to_html(para)
                    
                    if inner_html:
                        if info['is_list']:
                            current_content.append(f"<li>{inner_html}</li>")
                        elif style_type == 'heading':
                            current_content.append(f"<h{level}>{inner_html}</h{level}>")
                        elif style_type == 'quote':
                            current_content.append(f"<blockquote>{inner_html}</blockquote>")
                        else:
                            # Alinhamento
                            alignment = ""
                            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                                alignment = " style='text-align:center'"
                            elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                                alignment = " style='text-align:right'"
                            elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                                alignment = " style='text-align:justify'"
                            current_content.append(f"<p{alignment}>{inner_html}</p>")
                
            elif element.tag.endswith('tbl'):
                # Fecha listas antes de tabelas
                current_content.append(list_manager.close_all())
                
                table = tbl_map.get(element)
                if table:
                    current_content.append(self._table_to_html(table))
        
        # Fecha todas as listas abertas no final
        current_content.append(list_manager.close_all())
        
        if current_content or not sections:
            sections.append((
                current_title,
                '\n'.join(current_content) if current_content else "<p>(Conteúdo vazio)</p>",
                self._generate_smart_id(current_title, current_section_month)
            ))
        
        return sections
    
    def _document_to_single_entry(self) -> Tuple[str, str]:
        """
        Converte documento inteiro em uma única entrada, suportando listas.
        """
        title = self.filepath.stem
        content_parts = []
        list_manager = HTMLListManager(self._get_list_tag)
        
        # Otimização O(N): Mapeia element para parágrafo/tabela
        p_map = {p._element: p for p in self.document.paragraphs}
        tbl_map = {t._element: t for t in self.document.tables}
        
        for element in self.document.element.body:
            if element.tag.endswith('p'):
                para = p_map.get(element)
                
                if para:
                    info = self._get_paragraph_info(para)
                    style_type = info['type']
                    level = info['level']
                    
                    # Gerencia tags de lista
                    list_tags_html = list_manager.handle_paragraph(info)
                    if list_tags_html:
                        content_parts.append(list_tags_html)
                    
                    inner_html = self._paragraph_to_html(para)
                    
                    if inner_html:
                        if info['is_list']:
                            content_parts.append(f"<li>{inner_html}</li>")
                        elif style_type == 'heading':
                            content_parts.append(f"<h{level}>{inner_html}</h{level}>")
                        elif style_type == 'quote':
                            content_parts.append(f"<blockquote>{inner_html}</blockquote>")
                        else:
                            # Alinhamento
                            alignment = ""
                            if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                                alignment = " style='text-align:center'"
                            elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                                alignment = " style='text-align:right'"
                            elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                                alignment = " style='text-align:justify'"
                            content_parts.append(f"<p{alignment}>{inner_html}</p>")
            
            elif element.tag.endswith('tbl'):
                # Fecha listas antes de tabelas
                content_parts.append(list_manager.close_all())
                
                table = tbl_map.get(element)
                if table:
                    content_parts.append(self._table_to_html(table))
        
        # Fecha todas as listas abertas no final
        content_parts.append(list_manager.close_all())
        
        return (title, '\n'.join(content_parts))
    
    def convert(self, options: Optional[ConversionOptions] = None, **kwargs) -> MySwordJournal:
        """
        Converte o documento Word para Journal MySword.
        
        Metadados são extraídos automaticamente do documento Word:
        - Das propriedades do arquivo (Arquivo > Propriedades)
        - De linhas no início do documento (Autor:, Descrição:, Tags:, Abreviação:)
        - Do primeiro título (Heading 1) como título do Journal
        
        Args:
            options: Objeto ConversionOptions com todas as opções
            **kwargs: Argumentos alternativos (sobrescrevem os extraídos)
            
        Returns:
            Objeto MySwordJournal pronto para salvar
        """
        # Usa metadados extraídos do documento como base
        extracted = self._extracted_metadata
        
        # Processa opções (kwargs sobrescrevem metadados extraídos)
        if options is None:
            # Prioridade: kwargs > metadados extraídos > valores padrão
            abbr = kwargs.get('abbreviation') or extracted.get('abbreviation') or self.filepath.stem.replace(' ', '_')[:20]
            options = ConversionOptions(
                abbreviation=abbr,
                title=kwargs.get('title') or extracted.get('title') or self.filepath.stem,
                description=kwargs.get('description') or extracted.get('description') or "",
                author=kwargs.get('author') or extracted.get('author') or "",
                readonly=kwargs.get('readonly', kwargs.get('read_only', False)),
                split_by_heading=kwargs.get('split_by_heading', True),
                heading_level=kwargs.get('heading_level', 1),
                preserve_formatting=kwargs.get('preserve_formatting', True),
                convert_tables=kwargs.get('convert_tables', True),
                customcss=kwargs.get('customcss', kwargs.get('custom_css', ''))
            )
        
        # Cria o Journal
        journal = MySwordJournal(
            abbreviation=options.abbreviation,
            title=options.title,
            description=options.description,
            author=options.author,
            readonly=options.readonly,
            customcss=options.customcss
        )
        
        # Converte conteúdo
        if options.split_by_heading:
            sections = self._split_by_headings(options.heading_level)
            for title, content, entry_id in sections:
                journal.add_entry(title, content, id=entry_id)
        else:
            title, content = self._document_to_single_entry()
            journal.add_entry(title, content)
        
        return journal


def convert_word_to_journal(
    input_file: str,
    output_file: str,
    abbreviation: Optional[str] = None,
    title: str = "",
    description: str = "",
    author: str = "",
    split_by_heading: bool = True,
    heading_level: int = 1
) -> Path:
    """
    Função de conveniência para converter Word para Journal.
    
    Metadados são extraídos automaticamente do documento Word se não fornecidos:
    - Das propriedades do arquivo (Arquivo > Propriedades)
    - De linhas no início (Autor:, Descrição:, etc)
    
    Args:
        input_file: Caminho do arquivo .docx
        output_file: Caminho do arquivo .jor.mybible de saída
        abbreviation: Abreviação do Journal (extrai do documento se não fornecido)
        title: Título exibido no MySword (extrai do documento se não fornecido)
        description: Descrição do Journal (extrai do documento se não fornecido)
        author: Autor do Journal (extrai do documento se não fornecido)
        split_by_heading: Se True, divide por títulos
        heading_level: Nível do título para divisão
        
    Returns:
        Path do arquivo .jor.mybible criado
    """
    converter = WordToJournalConverter(input_file)
    
    # Passa apenas parâmetros que foram explicitamente fornecidos
    # Os vazios serão preenchidos com metadados extraídos
    kwargs = {
        'split_by_heading': split_by_heading,
        'heading_level': heading_level
    }
    
    # Só passa se não estiver vazio
    if abbreviation:
        kwargs['abbreviation'] = abbreviation
    if title:
        kwargs['title'] = title
    if description:
        kwargs['description'] = description
    if author:
        kwargs['author'] = author
    
    journal = converter.convert(**kwargs)
    
    return journal.save(output_file)


# CLI usando Click
if __name__ == "__main__":
    import click
    
    @click.command()
    @click.argument('input_file', type=click.Path(exists=True))
    @click.argument('output_file', type=click.Path())
    @click.option('--title', '-t', help='Título do Journal')
    @click.option('--description', '-d', help='Descrição do Journal')
    @click.option('--author', help='Autor do Journal')
    @click.option('--abbreviation', '-a', help='Abreviação (sem espaços)')
    @click.option('--no-split', is_flag=True, help='Não dividir por títulos')
    @click.option('--heading-level', '-h', default=1, type=int, help='Nível do heading para divisão')
    def main(input_file, output_file, title, description, author, abbreviation, no_split, heading_level):
        """
        Converte documento Word para Journal do MySword.
        
        Exemplo:
            python word_to_journal.py estudo.docx estudo.jor.mybible --title "Meus Estudos"
        """
        try:
            output_path = convert_word_to_journal(
                input_file=input_file,
                output_file=output_file,
                abbreviation=abbreviation,
                title=title or "",
                description=description or "",
                author=author or "",
                split_by_heading=not no_split,
                heading_level=heading_level
            )
            click.echo(f"✓ Journal criado: {output_path}")
            
        except Exception as e:
            click.echo(f"✗ Erro: {e}", err=True)
            raise SystemExit(1)
    
    main()
